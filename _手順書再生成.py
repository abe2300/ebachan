# -*- coding: utf-8 -*-
"""定例報告アプリ 運用手順書 (.docx) 生成スクリプト"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

THIS_DIR = os.path.dirname(os.path.abspath(__file__))
OUT_PATH = os.path.join(THIS_DIR, "定例報告_運用手順書.docx")

# テーマカラー（紺色ベース）
THEME_COLOR = "1E3A8A"
ACCENT_COLOR = "F59E0B"

doc = Document()

# ----- 既定スタイル -----
style = doc.styles["Normal"]
style.font.name = "游ゴシック"
style.font.size = Pt(10.5)
rPr = style.element.get_or_add_rPr()
rFonts = rPr.find(qn("w:rFonts"))
if rFonts is None:
    rFonts = OxmlElement("w:rFonts")
    rPr.append(rFonts)
rFonts.set(qn("w:eastAsia"), "游ゴシック")
rFonts.set(qn("w:ascii"), "游ゴシック")
rFonts.set(qn("w:hAnsi"), "游ゴシック")

# ページ余白
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)


def set_jp_font(run, name="游ゴシック", size=None, bold=None, color=None, mono=False):
    if mono:
        name = "Consolas"
    run.font.name = name
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name if not mono else "ＭＳ ゴシック")
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_jp_font(r, size=18, bold=True, color=RGBColor(0x1E, 0x3A, 0x8A))
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), THEME_COLOR)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_jp_font(r, size=13.5, bold=True, color=RGBColor(0x33, 0x33, 0x33))
    return p


def add_h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("■ " + text)
    set_jp_font(r, size=11.5, bold=True, color=RGBColor(0x1E, 0x3A, 0x8A))
    return p


def add_p(text, bold=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_jp_font(r, size=10.5, bold=bold, color=color)
    return p


def add_bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    set_jp_font(r, size=10.5)
    return p


def add_number(text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    set_jp_font(r, size=10.5)
    return p


def add_note(text, color="FFF8E1", border="F59E0B"):
    """注意ボックス"""
    t = doc.add_table(rows=1, cols=1)
    t.autofit = True
    cell = t.cell(0, 0)
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color)
    tcPr.append(shd)
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "8")
        b.set(qn("w:color"), border)
        tcBorders.append(b)
    tcPr.append(tcBorders)
    p = cell.paragraphs[0]
    for line in text.split("\n"):
        if p.runs:
            p = cell.add_paragraph()
        r = p.add_run(line)
        set_jp_font(r, size=10)
    doc.add_paragraph()


def add_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        p = cell.paragraphs[0]
        r = p.add_run(h)
        set_jp_font(r, size=10.5, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), THEME_COLOR)
        tcPr.append(shd)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            set_jp_font(r, size=10)
    doc.add_paragraph()


# =============================================
# 表紙
# =============================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(80)
r = p.add_run("定例報告アプリ")
set_jp_font(r, size=28, bold=True, color=RGBColor(0x1E, 0x3A, 0x8A))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("運用手順書")
set_jp_font(r, size=20, bold=True, color=RGBColor(0x33, 0x33, 0x33))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
r = p.add_run("毎日の経済ニュース・税制改正情報を確認するPWAアプリ")
set_jp_font(r, size=12, color=RGBColor(0x66, 0x66, 0x66))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(200)
r = p.add_run("作成日：2026年6月2日　／　最終更新：2026年6月5日（GitHub Pages移行）")
set_jp_font(r, size=11, color=RGBColor(0x99, 0x99, 0x99))

doc.add_page_break()

# =============================================
# 1. アプリの概要
# =============================================
add_h1("1. アプリの概要")
add_p("「定例報告アプリ」は、毎日の経済ニュース・税制改正情報を一覧できるPWA（プログレッシブウェブアプリ）です。PC・スマホ・タブレットから同じ画面を確認でき、複数人での情報共有も可能です。")

add_h3("主な機能")
add_bullet("主要経済ニュース 5件（各ニュースに株価への影響コメント付き）")
add_bullet("税制改正情報（会社員関連項目は詳細解説モーダル付き）")
add_bullet("税制改正の年度切替機能（過去年度の参照可、自動年度判定）")
add_bullet("日付の自動更新（今日の日付・報告データ更新日）")
add_bullet("印刷／PDF出力ボタン")
add_bullet("全文コピー機能")
add_bullet("報告ルール・情報ソースの参照")

add_h3("特徴")
add_bullet("ローカル起動：自分のPCで完結（オフライン対応）")
add_bullet("Web公開：GitHub Pages 自動デプロイ（push したら約1分で公開、完全無料・クレジット制限なし）")
add_bullet("PWA対応：スマホのホーム画面に追加してアプリのように使える")
add_bullet("ワンクリック公開：「GitHubに更新.bat」だけで全員に最新情報が配信される")

add_h3("公開URL（現行：GitHub Pages）")
add_p("https://abe2300.github.io/ebachan/")
add_p("※ 2026年6月5日にNetlifyのクレジット切れによりGitHub Pagesへ移行しました。", color=RGBColor(0x99, 0x99, 0x99))

add_h3("旧公開URL（停止中・参考）")
add_p("https://teirei-houkoku.netlify.app  （Netlifyクレジット限度超過のため一時停止）", color=RGBColor(0x99, 0x99, 0x99))

add_h3("GitHubリポジトリ")
add_p("https://github.com/abe2300/ebachan  （URL slug: ebachan、表示名: エバチャン）")

# =============================================
# 2. フォルダ構成
# =============================================
add_h1("2. フォルダ構成")
add_p("「クロード用 > 定例報告」フォルダに、以下のファイルが含まれています。")

add_table(
    ["ファイル名", "役割"],
    [
        ["index.html", "メイン画面（HTML本体）"],
        ["manifest.json", "PWAマニフェスト（アプリ情報）"],
        ["sw.js", "Service Worker（オフライン対応）"],
        ["icon-192.png / icon-512.png", "アプリアイコン"],
        ["server.py", "ローカルHTTPサーバー（ポート8767）"],
        ["定例報告を起動.bat", "ローカル起動（ワンクリック）"],
        ["定例報告サーバー停止.bat", "ローカルサーバー停止"],
        ["★ GitHubに更新.bat", "★ ワンクリックでgit push（毎日使う）"],
        ["朝の自動更新.bat", "毎朝7時の自動起動用（B案・タスクスケジューラ）"],
        ["タスクスケジューラ_設定.bat", "Windows タスクスケジューラ設定（B案セットアップ）"],
        ["タスクスケジューラ_解除.bat", "Windows タスクスケジューラ解除"],
        ["_自動更新の指示プロンプト.md", "/schedule復旧時のプロンプトテンプレート"],
        [".git / .gitignore", "Git リポジトリ管理ファイル"],
        ["README.md", "GitHubで表示される説明書き"],
        ["_手順書再生成.py", "本手順書を再生成するスクリプト"],
        ["公開用/、公開用を更新.bat", "旧Netlify方式（GitHub連携後は使用しない）"],
    ],
)
add_note("★マークが付いているファイルが日常的によく使うものです。\n旧Netlify方式の「公開用」フォルダはGitHub連携で不要になりましたが、念のため残してあります。")

# =============================================
# 3. ローカルで起動する方法
# =============================================
add_h1("3. ローカルで起動する方法（自分のPCで使う）")

add_h2("起動手順")
add_number("「定例報告」フォルダを開く")
add_number("「定例報告を起動.bat」をダブルクリック")
add_number("Microsoft Edge がアプリモード（URLバーなし）で自動起動")
add_number("画面に主要経済ニュース・税制改正情報が表示される")

add_note("💡 ヒント：BATファイルを右クリック→「ショートカットの作成」でデスクトップに置いておくと、毎朝ワンクリックで起動できます。")

add_h2("停止手順")
add_p("基本的にはブラウザを閉じるだけでOKです。サーバーを完全停止したい場合：")
add_number("「定例報告サーバー停止.bat」をダブルクリック")
add_number("バックグラウンドで動いているPythonサーバーが停止")

add_h2("起動の仕組み")
add_bullet("BATファイルが Python の HTTPサーバー（pythonw）をバックグラウンド起動")
add_bullet("サーバーは 127.0.0.1:8767 でローカル待機")
add_bullet("Microsoft Edge をアプリモードで開き、http://localhost:8767/index.html を表示")

# =============================================
# 4. Web公開・共有する方法（GitHub Pages 自動デプロイ）
# =============================================
add_h1("4. Web公開・共有する方法（GitHub Pages 自動デプロイ）")
add_p("GitHub Pages を使って、git push するだけで自動的にWebサイトが更新される仕組みを構築済みです。完全無料・クレジット制限なしで、Netlifyに依存しません。")

add_h2("4-1. 全体の仕組み（既に構築済み）")
add_p("以下の自動連携が既に完成しています。")
add_table(
    ["段階", "内容"],
    [
        ["①", "ローカルの index.html を編集"],
        ["②", "「GitHubに更新.bat」をダブルクリック"],
        ["③", "自動で git add → commit → push が実行される"],
        ["④", "GitHubのリポジトリ（abe2300/ebachan）に反映"],
        ["⑤", "GitHub Pages が自動検知してビルド・配信"],
        ["⑥", "約1分で https://abe2300.github.io/ebachan/ に公開"],
    ],
)

add_h2("4-2. 設定情報（既に完了済み）")
add_table(
    ["項目", "値"],
    [
        ["公開URL（現行）", "https://abe2300.github.io/ebachan/"],
        ["公開方式", "GitHub Pages（Settings → Pages → main / root）"],
        ["GitHubリポジトリ", "https://github.com/abe2300/ebachan"],
        ["デプロイブランチ", "main"],
        ["公開フォルダ", "/ (root)"],
        ["ビルド", "なし（静的ファイル直接配信）"],
        ["Git ユーザー名", "abe2300"],
        ["Git メールアドレス", "a161046@gmail.com"],
    ],
)

add_h2("4-3. ワンクリック公開の使い方")
add_number("Claude Code で「定例報告」と入力し、index.html を更新してもらう")
add_number("デスクトップまたはフォルダ内の「GitHubに更新.bat」をダブルクリック")
add_number("黒い画面で進行状況が表示される")
add_number("「Push成功!」メッセージが出れば完了")
add_number("約1分以内に公開URLで最新版を確認できる")

add_note("💡 ヒント：「GitHubに更新.bat」を右クリック→「ショートカットの作成」でデスクトップに置いておくと便利です。")

add_h2("4-4. 共有方法")
add_bullet("公開URLをLINE・メール・Teams・社内チャットで送るだけ")
add_bullet("受け取った人はブラウザでURLを開くだけで閲覧可能")
add_bullet("スマホでも同じURLでアクセス可能（PWA化も可）")
add_bullet("常に最新版が表示される（受け取り側で更新作業不要）")

add_h2("4-5. セキュリティと制限")
add_bullet("リポジトリは Public（公開）設定 — コードは誰でも閲覧可能")
add_bullet("ただし業務上の機密情報は含まれていない（ニュース・税制情報のみ）")
add_bullet("URLを知っている人全員がアプリにアクセス可能")
add_bullet("GitHub Pages：完全無料、月間100GB帯域、ビルド時間制限なし（静的のため）")
add_bullet("GitHub無料枠：Public無制限、Privateも一定枠まで無料")

add_h2("4-6. Netlify運用からの移行について（経緯）")
add_p("2026年6月5日まではNetlify（https://teirei-houkoku.netlify.app）で運用していましたが、Netlifyのアカウントクレジット限度額超過により本番展開がブロックされたため、GitHub Pagesへ移行しました。GitHub Pagesはクレジット制限がないため、永続的に無料運用できます。")
add_bullet("旧URL： https://teirei-houkoku.netlify.app （停止中、月次リセットで復活する可能性あり）")
add_bullet("新URL： https://abe2300.github.io/ebachan/ （現行）")
add_bullet("「公開用」フォルダと「公開用を更新.bat」は旧Netlify Drop方式の名残で、現在は使いません。")

# =============================================
# 5. スマホで使う方法
# =============================================
add_h1("5. スマホで使う方法（PWA化）")
add_p("公開されたURLにスマホでアクセスすると、アプリのようにホーム画面に追加できます。")

add_h2("iPhone (Safari)")
add_number("Safari で公開URL（https://abe2300.github.io/ebachan/）を開く")
add_number("画面下の「共有」ボタン（□に↑のマーク）をタップ")
add_number("メニューから「ホーム画面に追加」を選択")
add_number("名前を確認して「追加」をタップ")
add_number("ホーム画面に「定例報告」のアイコンが追加される")

add_h2("Android (Chrome)")
add_number("Chromeで公開URL（https://abe2300.github.io/ebachan/）を開く")
add_number("右上の「︙」メニューをタップ")
add_number("「ホーム画面に追加」を選択")
add_number("名前を確認して「追加」をタップ")

add_note("💡 PWA化のメリット：オフラインでも閲覧可、起動が早い、ブラウザのURLバーが表示されない（アプリ風）")

add_h3("旧PWA（Netlify版）からの移行")
add_p("以前 https://teirei-houkoku.netlify.app からPWAをインストールしていた場合は、以下の手順で再インストールしてください：")
add_number("ホーム画面の旧「定例報告」アイコンを長押し → 削除")
add_number("新公開URL（https://abe2300.github.io/ebachan/）にスマホブラウザでアクセス")
add_number("上記手順で再度「ホーム画面に追加」")
add_note("⚠️ URLが変わったため、PWAは別アプリ扱いになります。再インストール必須です。")

add_h2("5-3. スマホでの編集について（重要）")
add_p("スマホはClaude.aiのchat機能のみのため、ファイルを直接編集することができません。スマホは閲覧専用と考えてください。")
add_table(
    ["環境", "閲覧", "編集"],
    [
        ["PC + Claude Code", "○", "○ ファイル直接編集可"],
        ["スマホ + claude.ai", "○", "△ チャットで生成→GitHub Web手動編集が必要"],
        ["スマホ + ブラウザのみ", "○ アプリ風に閲覧", "× 不可"],
    ],
)
add_note("📱 スマホで更新したい場合：claude.ai でニュース取得 → GitHub.com で index.html を手動編集 → コミット → GitHub Pages自動デプロイ。ただし所要時間10～15分とPC運用より大変なので、PC運用がおすすめです。")

# =============================================
# 6. 日々の運用フロー
# =============================================
add_h1("6. 日々の運用フロー")

add_h2("6-1. 毎朝の流れ（C案：ワンクリック公開）")
add_number("Claude Code を開く")
add_number("「定例報告」と入力 → Enter")
add_number("Claudeが最新の経済ニュース5件＋株価影響コメントを生成")
add_number("Claudeが index.html を直接更新（日付・ニュース内容）")
add_number("「GitHubに更新.bat」をダブルクリック")
add_number("「Push成功!」が表示されれば完了 — 約1分以内に公開URL（https://abe2300.github.io/ebachan/）に反映")
add_note("所要時間：1～2分。スマホ・他人もURLでアクセスすれば最新版を閲覧できます。")

add_h2("6-2. ニュース項目を変更する場合")
add_bullet("Claudeに「○○のニュースに差し替えて」と依頼")
add_bullet("Claudeが index.html を直接編集")
add_bullet("更新後、「GitHubに更新.bat」をダブルクリックでpush公開")

add_h2("6-3. 日付表示の見方")
add_p("画面ヘッダーに2種類の日付が表示されます。役割を理解しておくと、いつのデータかが分かりやすくなります。")
add_table(
    ["表示項目", "意味", "更新方法"],
    [
        ["大きな日付（例：2026年6月4日）", "アプリを開いた日（PCの今日の日付）", "毎日自動更新"],
        ["📌 報告データ更新日", "ニュース・税制情報の最終更新日", "Claudeに依頼して更新時のみ変更"],
        ["フッターの報告日", "アプリを開いた日（PCの今日の日付）", "毎日自動更新"],
    ],
)
add_note("💡 大きな日付と「報告データ更新日」が違う場合 → ニュース内容が古い可能性があります。Claudeに最新版への更新を依頼してください。")

add_h2("6-4. 税制改正の年度切替について")
add_p("税制改正情報セクションには、過去年度のデータも保管されており、ドロップダウンから選択できます。")

add_h3("年度バッジの自動判定")
add_table(
    ["バッジ", "色", "意味"],
    [
        ["最新", "緑", "現在の年度（4月～翌3月）— デフォルト表示"],
        ["前年度", "グレー", "1年前の年度"],
        ["過去", "薄グレー", "2年以上前の年度"],
        ["予定", "青", "未来の改正（早期発表時）"],
        ["更新待ち", "オレンジ", "今年度のデータが未登録（要更新）"],
    ],
)
add_note("4月1日になると、自動で「最新」バッジが次の年度に切り替わります。手動操作は不要です。")

add_h3("過去年度の参照方法")
add_number("税制改正情報セクションのドロップダウンをクリック")
add_number("見たい年度を選択（令和8年度 / 令和7年度 / 令和6年度 など）")
add_number("テーブルが瞬時に切り替わる")
add_number("詳細ボタンがある項目はクリックでモーダル表示")

add_h2("6-5. 新年度データを追加する方法")
add_p("毎年12月～1月頃に税制改正大綱が発表されるタイミングで、新年度のデータを追加してください。")

add_h3("Claudeへの依頼方法")
add_bullet("例：「令和9年度（2027年度）の税制改正データを追加して」")
add_bullet("Claudeが index.html の taxReformData に自動追加")
add_bullet("ドロップダウンに自動で表示される")
add_bullet("バッジは自動判定（西暦年から計算）なので手動設定不要")

add_h3("追加後の確認")
add_number("アプリを開き直す（または F5 で再読込）")
add_number("新年度がドロップダウンの最上段に表示されているか確認")
add_number("「最新」バッジが付いているか確認")
add_number("「GitHubに更新.bat」→ GitHub Pagesに公開")

add_h2("6-6. データ未追加の警告が出たら")
add_p("4月以降、「⚠️ 令和X年度のデータ未追加」というオレンジの警告が表示されることがあります。")
add_bullet("原因：新年度に入ったが、まだ新年度のデータが追加されていない")
add_bullet("対処：Claudeに「最新年度の税制改正データを追加して」と依頼")
add_bullet("追加後は警告が自動で消えます")

# =============================================
# 7. 自動化オプション（朝7時自動起動）
# =============================================
add_h1("7. 自動化オプション（毎朝7時の自動起動）")
add_p("通常はC案（ワンクリック公開）で十分ですが、毎朝7時に自動的に作業を始めたい場合は以下の方法があります。")

add_h2("7-1. B案：Windows タスクスケジューラ（PC起動中の自動化）")
add_p("Windows標準機能のタスクスケジューラを使って、毎朝7時に「朝の自動更新.bat」を起動させます。PC が起動中である必要があります。")

add_h3("初回セットアップ（1回だけ）")
add_number("「タスクスケジューラ_設定.bat」をダブルクリック")
add_number("確認メッセージで「Y」を押す")
add_number("「設定完了!」と表示されれば設定完了")
add_number("Windowsのタスクスケジューラに「定例報告_朝の自動起動」が登録される")

add_h3("毎朝7時の動作")
add_number("「朝の自動更新.bat」が自動起動")
add_number("デスクトップに通知が表示される")
add_number("Claude Code が自動起動（ブラウザ版の場合あり）")
add_number("ユーザーが「定例報告」と入力 → HTML 更新")
add_number("黒い画面で Enter キーを押す")
add_number("自動で git push → GitHub Pages公開")

add_h3("解除したい場合")
add_bullet("「タスクスケジューラ_解除.bat」を実行")
add_bullet("確認メッセージで「Y」を押すと解除される")

add_note("⚠️ 注意：B案はPCが起動している必要があります。スリープ・シャットダウンしているとタスクは動きません。確実に動かしたい場合は7時前にPC起動の習慣化を。")

add_h2("7-2. 将来：Claude Code /schedule（クラウド自動化）")
add_p("Claude Pro プランの /schedule 機能を使うと、クラウド側で自動実行されるため、PC起動不要で完全自動化が可能です（現在は接続不可で待機中）。")

add_h3("/schedule が使えるようになったら")
add_number("「タスクスケジューラ_解除.bat」で B案を解除")
add_number("Claude Code で「/schedule」を実行")
add_number("「_自動更新の指示プロンプト.md」の内容を /schedule に登録")
add_number("毎朝7時にクラウド側で自動実行")
add_number("PC が起動していなくても更新される")

add_h3("/schedule に渡すプロンプトの保管場所")
add_bullet("「_自動更新の指示プロンプト.md」に詳細を記載済み")
add_bullet("Claudeへの依頼内容、git push 手順まで網羅")

# =============================================
# 8. トラブルシューティング
# =============================================
add_h1("8. トラブルシューティング")

add_table(
    ["症状", "原因と対処"],
    [
        ["BATを実行しても起動しない", "Pythonがインストールされていない。python.org からインストール。"],
        ["ポート8767が使用中エラー", "「定例報告サーバー停止.bat」を実行してから再起動。"],
        ["Edge が起動しない", "別のブラウザ（Chrome）でも可。手動で http://localhost:8767/index.html を開く。"],
        ["GitHubに更新.batでpush失敗", "ネット接続を確認。git の認証が切れた場合はブラウザでGitHub再ログイン。"],
        ["GitHub Pagesに公開が反映されない", "Settings → Pages でデプロイ状況を確認。Actions タブで「pages-build-deployment」が成功しているか確認。約1分かかります。"],
        ["GitHub Pagesが404になる", "Settings → Pages で Source が「main / root」になっているか確認。初回有効化後、数分待つ必要があります。"],
        ["旧Netlify URLが繋がらない", "クレジット切れで停止中。新URL https://abe2300.github.io/ebachan/ を使用してください。"],
        ["スマホPWAで古い画面が表示される", "Service Worker のキャッシュ。sw.jsのCACHE名を v6→v7 などに変更してデプロイ。それでもダメな場合はPWA削除→再インストール。"],
        ["PWAアイコンが古い", "スマホのキャッシュを削除→再度ホーム画面に追加。"],
        ["旧PWA（Netlify版）がエラー", "URLが変わっているため使えません。新URL（GitHub Pages版）でPWAを再インストールしてください。"],
        ["日付が今日に更新されない", "JSが動いていない可能性。F5で再読込。年が変わっても自動更新されます。"],
        ["税制改正が古い年度のまま", "4月1日に自動切替されない場合、ブラウザのリロード（Ctrl+Shift+R）を試す。"],
        ["「データ未追加」警告が出る", "新年度のデータが未登録。Claudeに「最新年度の税制改正データを追加して」と依頼。"],
        ["タスクスケジューラが動かない", "PC が起動していない/スリープしている可能性。タスクスケジューラのアプリで実行履歴を確認。"],
        ["/schedule が「接続できない」と表示", "Anthropic側の一時的なメンテナンスまたはアカウント側の機能制限。後日再試行。"],
        ["手順書を作り直したい", "「_手順書再生成.py」をダブルクリックで再生成。"],
    ],
)

# =============================================
# 9. 手順書の再生成
# =============================================
add_h1("9. 手順書の再生成")
add_p("本手順書は Python スクリプトで自動生成しています。内容を変更したい場合：")
add_number("「_手順書再生成.py」をテキストエディタで開く")
add_number("修正したい箇所を編集")
add_number("ファイルをダブルクリックで実行")
add_number("「定例報告_運用手順書.docx」が上書き生成される")

add_note("必要なPythonライブラリ：python-docx\nインストール：pip install python-docx")

# =============================================
# 保存
# =============================================
doc.save(OUT_PATH)
print(f"手順書を生成しました：{OUT_PATH}")
